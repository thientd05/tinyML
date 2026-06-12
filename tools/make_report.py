"""Generate a DOCX progress report for the ECG-on-ESP32 project.

Reads results/summary.csv (+ the v1 backup summary_v1_before_clean.csv for the
before/after-cleaning comparison), results/winners.json (auto-selected winners, so the
report never hard-codes them), results/cleaning_stats.json, and the three plots in
results/, then writes results/BaoCao_TinyML_ECG.docx. Run from repo root:
    ./env/bin/python3.10 tools/make_report.py
"""
import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
RESULTS = ROOT / "results"
SUMMARY = RESULTS / "summary.csv"
SUMMARY_V1 = RESULTS / "summary_v1_before_clean.csv"
WINNERS = RESULTS / "winners.json"
CLEAN_STATS = RESULTS / "cleaning_stats.json"
OUT = RESULTS / "BaoCao_TinyML_ECG.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x59, 0x59, 0x59)

FAM_LABEL = {"rf": "RF", "svm": "SVM", "cnn": "CNN", "lstm": "LSTM"}


def load_rows(path=SUMMARY):
    if not path.exists():
        return []
    with open(path, newline="") as f:
        return list(csv.DictReader(f))


def load_winners():
    """Auto-selected winners written by analyze.py (never hard-coded here)."""
    d = json.loads(WINNERS.read_text())
    return d["within"], (tuple(d["cross"]) if d.get("cross") else None)


def load_clean_stats():
    return json.loads(CLEAN_STATS.read_text()) if CLEAN_STATS.exists() else None


def fnum(x, n=3):
    try:
        return f"{float(x):.{n}f}"
    except (TypeError, ValueError):
        return str(x)


# ---------- document helpers ----------
def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT
    return h


def para(doc, text, italic=False, size=11, color=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)
    return p


def add_image(doc, name, caption, width=6.3):
    doc.add_picture(str(RESULTS / name), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = GREY
    cap.paragraph_format.space_after = Pt(12)


def style_header_row(row):
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell._tc.get_or_add_tcPr().append(_shade("1F4E79"))


def _shade(hexcolor):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    return sh


def set_cell(cell, text, size=8.5, bold=False, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if fill:
        cell._tc.get_or_add_tcPr().append(_shade(fill))


# ---------- build report ----------
def section_data_cleaning(doc, rows, by_key, v1_by_key, within, clean):
    """Section 2: dataset + the (legitimate, no-leakage) cleaning applied before training,
    with a before/after table. All numbers come from cleaning_stats.json / the two CSVs."""
    heading(doc, "2. Dữ liệu và làm sạch dữ liệu", 1)
    para(doc,
         "Dữ liệu là MIT-BIH Arrhythmia DB (PhysioNet, 360 Hz, 2 chuyển đạo). Mỗi mẫu là "
         "một nhịp (±100 mẫu quanh đỉnh R, ~555 ms) trên chuyển đạo MLII; nhãn AAMI nhị "
         "phân (N,L,R,e,j → Normal; còn lại → Abnormal). Chia theo bệnh nhân (de Chazal: "
         "DS1 train / DS2 test), loại 4 bản ghi tạo nhịp 102/104/107/217. Lớp mất cân bằng "
         "~11% Abnormal.")
    para(doc,
         "Trước khi train, dữ liệu được làm sạch ở những phần CÓ THỂ sửa mà không “ăn gian”. "
         "Phân biệt rõ: độ khó do mất cân bằng lớp và do chia theo bệnh nhân là độ khó THẬT "
         "của bài toán — giữ nguyên. Chỉ làm sạch phần chồng lấn do lỗi pipeline và nhiễu. "
         "Nguyên tắc chống rò rỉ: mọi phép làm sạch là label-free (chỉ dựa trên chuyển đạo, "
         "cấu trúc annotation và hình dạng tín hiệu, KHÔNG dùng nhãn Normal/Abnormal), áp "
         "y hệt cho cả train / validation / test, và chuẩn hóa z-score theo từng bản ghi.",
         italic=True, size=10.5, color=GREY)

    sp = clean["splits"] if clean else {}
    da = sum(sp.get(k, {}).get("dropped_artifact", 0) for k in sp)
    by = [0, 0]
    for k in sp:
        bl = sp.get(k, {}).get("dropped_artifact_by_label", [0, 0])
        by[0] += bl[0]; by[1] += bl[1]
    test_drop = sp.get("test", {}).get("dropped_artifact", 0)

    bullet(doc, "bản ghi 114 để chuyển đạo V5 ở kênh 0 (15/16 bản ghi khác là MLII ở kênh "
                "0). Trước đây nạp mù kênh 0 → 114 bị đánh giá trên sai chuyển đạo. Sửa: chọn "
                "kênh MLII THEO TÊN. Là lỗi nạp dữ liệu, sửa là đúng đắn.",
           "1. Sai chuyển đạo (record 114): ")
    bullet(doc, "khoảng RR trước đây tính tới annotation liền kề trong mảng đầy đủ — có thể "
                "là dấu không-phải-nhịp (‘+’ đổi nhịp điệu, ‘~’…), làm RR sai ở ~4.2% nhịp "
                "(DS2 ~4.7%). Sửa: chỉ tính RR giữa các NHỊP THẬT liên tiếp. Ảnh hưởng đặc "
                "trưng RF/SVM.", "2. Nhiễm khoảng RR (~4.2%): ")
    bullet(doc, "cửa sổ nhịp trước đây cắt cứng quanh mẫu annotation (lệch tới ~11–13 mẫu ở "
                "p95–p99). Sửa: căn lại cửa sổ vào cực trị |tín hiệu| trong ±15 mẫu để hình "
                "nhịp sắc nét, label-free.", "3. Căn lại đỉnh R: ")
    bullet(doc, f"loại nhịp VẬT LÝ không dùng được (flatline/lead-off + bão hòa/clipping) "
                f"theo tiêu chí tín hiệu khách quan — KHÔNG loại theo biên độ (tránh thiên "
                f"lệch các nhịp bất thường biên độ lớn). Thực tế chỉ loại {da} nhịp toàn bộ "
                f"(Normal={by[0]}, Abnormal={by[1]} → không thiên lệch lớp), và DS2 mất "
                f"{test_drop} nhịp — bằng chứng KHÔNG cherry-pick tập test.",
           "4. Lọc nhịp nhiễu (signal-based): ")

    # before/after table for each family's (new) winner
    heading(doc, "2.1. Tác động: precision @ recall 0.95 trước/sau làm sạch", 2)
    t = doc.add_table(rows=1, cols=5)
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, name in enumerate(["Mô hình thắng (mỗi họ)", "Trước (v1)", "Sau (v2)", "Δ", "rec. triển khai (v2)"]):
        set_cell(t.rows[0].cells[j], name, size=9, bold=True)
    style_header_row(t.rows[0])
    for fam in ["rf", "svm", "cnn", "lstm"]:
        size = within.get(fam)
        if not size:
            continue
        v2 = by_key.get((fam, size)); v1 = v1_by_key.get((fam, size))
        p2 = float(v2["precision_op"]); s2 = float(v2.get("precision_op_std") or 0.0)
        p1 = float(v1["precision_op"]) if v1 else float("nan")
        d = p2 - p1
        cells = t.add_row().cells
        vals = [f"{FAM_LABEL[fam]}/{size}", fnum(p1, 3) if v1 else "—",
                f"{p2:.3f}±{s2:.2f}", f"{d:+.3f}" if v1 else "—", fnum(v2["recall_deploy"], 3)]
        for j, v in enumerate(vals):
            set_cell(cells[j], v, size=8.5)
    para(doc, "“Trước (v1)” = cùng cấu hình train trên dữ liệu CHƯA làm sạch (single-run); "
              "“Sau (v2)” = trung bình ± std trên nhiều seed sau làm sạch. Phần lớn Δ nằm trong "
              "dải nhiễu seed của CNN/LSTM — cho thấy tác động chính của làm sạch là sửa lỗi "
              "(lead/RR) chứ không phải thổi phồng chất lượng.", italic=True, size=9.5,
         color=GREY, space_after=10)


def ensure_beat_figures():
    """Section 6.2 embeds two illustrative beat figures rendered from MIT-BIH by
    tools/make_beat_figures.py. Regenerate them if missing so this stays one-command
    (needs the dataset under dataset/; skipped gracefully if unavailable)."""
    needed = [RESULTS / "beat_easy_vs_hard.png", RESULTS / "beat_morphology_gallery.png"]
    if all(p.exists() for p in needed):
        return
    try:
        import sys
        sys.path.insert(0, str(ROOT))
        from tools import make_beat_figures as mbf
        mbf.gallery()
        mbf.easy_vs_hard()
    except Exception as e:  # dataset missing / wfdb absent — leave existing pngs or fail loudly later
        print(f"[make_report] WARN: could not render beat figures ({e}); "
              "run tools/make_beat_figures.py manually.")


def main():
    ensure_beat_figures()
    rows = load_rows()
    by_key = {(r["family"], r["size"]): r for r in rows}
    v1_by_key = {(r["family"], r["size"]): r for r in load_rows(SUMMARY_V1)}
    within, cross = load_winners()
    clean = load_clean_stats()

    # cross-family winner row + derived numbers (parameterized; never hard-coded)
    cfam, csize = cross
    crow = by_key[(cfam, csize)]
    cprec, clat = float(crow["precision_op"]), float(crow["latency_ms"])
    cprec_std = float(crow.get("precision_op_std") or 0.0)
    cflash, crecdep = float(crow["flash_kb"]), float(crow["recall_deploy"])
    win_precs = [float(by_key[(f, s)]["precision_op"]) for f, s in within.items()]
    pmin, pmax = min(win_precs), max(win_precs)
    n_total = len(rows)
    n_seeds = int(float(crow.get("n_seeds") or 1))

    def pm(r):  # "mean±std" for precision_op
        return f"{float(r['precision_op']):.3f}±{float(r.get('precision_op_std') or 0.0):.2f}"

    def drec(fam):  # DS2 recall at the DS1-val-tuned (deployment) threshold
        return float(by_key[(fam, within[fam])]["recall_deploy"]) if fam in within else float("nan")
    classic_rec = sorted(x for x in (drec("rf"), drec("svm")) if x == x)
    nn_rec = sorted(x for x in (drec("cnn"), drec("lstm")) if x == x)
    rec_classic = (f"~{classic_rec[0]:.2f}–{classic_rec[-1]:.2f}" if classic_rec else "n/a")
    rec_nn = (f"~{nn_rec[0]:.2f}–{nn_rec[-1]:.2f}" if len(nn_rec) > 1 else
              (f"~{nn_rec[0]:.2f}" if nn_rec else "n/a"))

    doc = Document()
    # base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ===== Title =====
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("BÁO CÁO TIẾN ĐỘ\nPhát hiện bất thường nhịp tim ECG bằng TinyML trên ESP32")
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = ACCENT
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(
        "So sánh 4 họ mô hình (Random Forest · SVM · 1D-CNN · LSTM) "
        "trên MIT-BIH Arrhythmia\nbằng phương pháp lựa chọn có cơ sở "
        "(capacity sweep + operating point + feasibility gate)"
    )
    rs.italic = True
    rs.font.size = Pt(11)
    rs.font.color.rgb = GREY
    doc.add_paragraph()

    # ===== 1. Tóm tắt =====
    heading(doc, "1. Tóm tắt kết quả", 1)
    para(
        doc,
        "Dự án xây dựng và so sánh 4 họ mô hình học máy cho bài toán phát hiện nhịp tim "
        "bất thường (Normal vs Abnormal) chạy thật trên vi điều khiển ESP32-WROOM-32. "
        "Việc lựa chọn mô hình dựa trên một khung có cơ sở, mã hóa đúng 3 yêu cầu của bài "
        "toán lâm sàng: vừa bộ nhớ ESP32, phát hiện tức thời, và tối thiểu bỏ sót ca bất thường.",
    )
    bullet(doc, f"Đã quét {n_total} cấu hình (4 họ × 6 điểm dung lượng), huấn luyện trên PC "
                f"và triển khai cả {n_total} lên ESP32.", f"{n_total} mô hình: ")
    bullet(doc, "dữ liệu được làm sạch (sửa chuyển đạo record 114, tính RR giữa các nhịp "
                "thật, căn lại đỉnh R, lọc nhịp nhiễu) theo cách không rò rỉ nhãn TRƯỚC khi "
                "train — chi tiết ở mục 2.", "Làm sạch dữ liệu: ")
    bullet(doc, f"tất cả {n_total} mô hình khớp bit-for-bit (parity = 1.0000) giữa ESP32 và "
                "PC (sklearn/PyTorch) → số liệu PC dùng được làm chuẩn.", "Đúng đắn: ")
    bullet(doc, f"mỗi cấu hình được train lại trên {n_seeds} seed; báo cáo trung bình ± độ "
                "lệch chuẩn và CHỌN winner theo trung bình — vì precision @ recall 0.95 của "
                "CNN/LSTM dao động ~±0.05 chỉ do seed, single-run dễ cho winner ‘ảo’.",
           f"So sánh bền vững ({n_seeds} seed): ")
    bullet(doc, f"{FAM_LABEL[cfam]} “{csize}” — precision {cprec:.3f}±{cprec_std:.2f} @ recall "
                f"0.95, độ trễ {clat:.1f} ms/nhịp, flash {cflash:.1f} KB. Đáp ứng cả 3 ràng buộc.",
           "Mô hình thắng cuộc (theo trung bình): ")
    bullet(doc, f"“to hơn” không đồng nghĩa “tốt hơn”; ngưỡng quyết định cần tinh chỉnh theo "
                f"độ nhạy mục tiêu; và độ khó thật của bài toán ở mức precision "
                f"{pmin:.2f}–{pmax:.2f} tại recall 0.95.", "3 phát hiện chính: ")

    # ===== 2. Dữ liệu & làm sạch =====
    section_data_cleaning(doc, rows, by_key, v1_by_key, within, clean)

    # ===== 3. Phương pháp =====
    heading(doc, "3. Phương pháp lựa chọn mô hình", 1)
    para(
        doc,
        "Bài toán đặt ra: thiết kế model thuộc 4 họ ML ở nhiều kích thước, đo hiệu năng – tốc "
        "độ – khả năng tải – thông số trên ESP32, rồi so sánh trong cùng một họ và giữa các họ. "
        "Ngoài ràng buộc vừa bộ nhớ ESP32, tiêu chí tối ưu là phát hiện đủ tức thời và bỏ sót "
        "ca bất thường ít nhất (vì bỏ sót là nguy hiểm nhất).",
    )
    para(
        doc,
        "Việc chọn model được phát biểu như một bài toán tối ưu có ràng buộc, tách rõ hai trục: "
        "dung lượng model (capacity) cho ta năng lực, còn ngưỡng quyết định (operating point) "
        "cho ta cách vận hành. Cụ thể gồm 3 thành phần:",
    )
    bullet(doc, "mỗi họ quét ~6 điểm trên một “núm” dung lượng đơn điệu (RF: số node "
                "cây; SVM: số support vector; CNN/LSTM: số phép nhân-cộng MACs). Nhờ đó vẽ được "
                "đường cong chất-lượng-vs-chi-phí thay vì 4 điểm rời rạc.", "Capacity sweep: ")
    bullet(doc, "thay vì ngưỡng 0.5, ta chọn ngưỡng đạt recall ≥ 0.95 trên lớp bất thường, "
                "tinh chỉnh trên tập DS1-val tách riêng theo bệnh nhân (record 207/209/215), rồi "
                "báo cáo trên tập kiểm thử DS2. Đây chính là mã hóa tiêu chí “bỏ sót ít nhất”.",
           "Operating point (recall ≥ 0.95): ")
    bullet(doc, "một cấu hình chỉ “hợp lệ” khi đồng thời: latency end-to-end ≤ 100 ms "
                "(đo thật trên ESP32, RF/SVM cộng thêm 0.94 ms chi phí trích đặc trưng), flash "
                "≤ 1 MB, và RAM trong ngân sách. Trong vùng hợp lệ, chọn model có precision cao "
                "nhất tại recall 0.95 (tức ít báo động giả nhất).", "Feasibility gate: ")
    para(
        doc,
        "Ngân sách 100 ms tương ứng ~3 lần biên an toàn so với nhịp tim nhanh ~200 bpm (cách "
        "nhau ~300 ms) — đủ “tức thời”. Dùng PR-AUC (average precision) thay ROC-AUC "
        "vì dữ liệu mất cân bằng (~11% bất thường). Đánh giá chất lượng lấy từ toàn bộ DS2 trên "
        "PC (chính xác hơn), ESP32 chỉ đo độ trễ và RAM — hợp lệ nhờ parity = 1.0000.",
        italic=True, size=10.5, color=GREY,
    )
    bullet(doc, f"vì precision @ recall 0.95 của mạng raw-beat dao động mạnh theo seed khởi tạo "
                f"(~±0.05), mỗi cấu hình được train lại trên {n_seeds} seed; báo cáo và xếp hạng "
                "theo TRUNG BÌNH (kèm std). Tránh đúng cái bẫy single-run từng tạo ra một "
                "‘winner’ không tái lập được.", "Độ bền theo seed: ")

    # ===== 4. Kết quả: bảng tổng hợp =====
    heading(doc, "4. Kết quả", 1)
    heading(doc, f"4.1. Bảng tổng hợp {n_total} cấu hình (đánh giá trên DS2 — bệnh nhân chưa từng thấy)", 2)

    cols = [
        ("Mô hình", lambda r: f"{r['family']}/{r['size']}"),
        ("MACs", lambda r: f"{int(float(r['macs'])):,}"),
        ("Flash KB", lambda r: fnum(r["flash_kb"], 1)),
        ("Latency ms", lambda r: fnum(r["latency_ms"], 1)),
        ("Recall@op", lambda r: fnum(r["recall_op"], 2)),
        ("Prec@op (±std)", lambda r: pm(r)),
        ("FPR", lambda r: fnum(r["fpr_op"], 2)),
        ("PR-AUC", lambda r: fnum(r["pr_auc"], 3)),
        ("Hợp lệ", lambda r: "✓" if r["feasible"] == "True" else "✗"),
    ]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, (name, _) in enumerate(cols):
        set_cell(table.rows[0].cells[j], name, size=9, bold=True)
    style_header_row(table.rows[0])

    fam_order = ["rf", "svm", "cnn", "lstm"]
    rows_sorted = sorted(rows, key=lambda r: (fam_order.index(r["family"]), float(r["macs"])))
    for r in rows_sorted:
        cells = table.add_row().cells
        is_win = within.get(r["family"]) == r["size"]
        infeasible = r["feasible"] != "True"
        fill = "DDEBF7" if is_win else ("F2DCDB" if infeasible else None)
        for j, (_, fn) in enumerate(cols):
            set_cell(cells[j], fn(r), size=8.5, bold=is_win, fill=fill)
    para(doc, f"Hàng xanh = mô hình thắng trong họ (theo precision trung bình); hàng hồng = vi "
              f"phạm ràng buộc (latency hoặc flash). “op” = tại ngưỡng đạt recall 0.95. Mỗi "
              f"precision là trung bình ± std trên {n_seeds} seed.", italic=True, size=9.5,
         color=GREY, space_after=10)

    # ===== 4.2 within-family =====
    heading(doc, "4.2. So sánh trong từng họ — đường chất lượng vs chi phí", 2)
    para(doc, "Mỗi đường là một họ; trục hoành là chi phí tính toán (log). Ta tìm điểm “knee” "
              "— nơi tăng chi phí không còn đổi lấy chất lượng.")
    add_image(doc, "within_family_quality_vs_cost.png",
              "Hình 1. Trong-họ: precision @ recall 0.95 (trái) và PR-AUC (phải) theo chi phí. "
              "RF gần như phẳng; SVM/CNN/LSTM có đỉnh ở model nhỏ rồi đi xuống.")
    bullet(doc, "tăng dung lượng gần như không đổi precision — dừng ở model nhỏ là hợp lý.", "RF: ")
    bullet(doc, f"đạt đỉnh ở model nhỏ ({FAM_LABEL['cnn']} “{within.get('cnn','?')}”) rồi giảm "
                "khi to thêm — cho thấy “to hơn không tốt hơn”.", "CNN: ")
    bullet(doc, "PR-AUC (phải) độc lập ngưỡng cho cùng kết luận; LSTM kém ổn định nhất.", "Chốt: ")

    # ===== 4.3 cross-family pareto =====
    heading(doc, "4.3. So sánh giữa các họ — Pareto front", 2)
    para(doc, "Mỗi điểm là một cấu hình; vùng trái đường 100 ms là hợp lệ về độ trễ. Vòng tròn "
              "đen = model thắng của mỗi họ. Càng lên cao càng ít báo động giả.")
    add_image(doc, "cross_family_pareto.png",
              "Hình 2. Pareto giữa các họ: precision @ recall 0.95 theo độ trễ end-to-end (log). "
              f"{FAM_LABEL[cfam]}/{csize} nằm cao nhất trong vùng hợp lệ.")

    # ===== 4.4 PR curves =====
    heading(doc, "4.4. Đường Precision–Recall của model tốt nhất mỗi họ", 2)
    add_image(doc, "pr_curves.png",
              "Hình 3. Đường PR trên DS2; đường đứng đỏ là mốc recall 0.95 — điểm vận hành "
              "thực tế. Tại đó precision của mọi họ đều thấp, phản ánh độ khó thật của bài toán.")

    # ===== 4.5 winners table =====
    heading(doc, "4.5. Mô hình thắng cuộc", 2)
    wt = doc.add_table(rows=1, cols=6)
    wt.style = "Table Grid"
    wt.alignment = WD_TABLE_ALIGNMENT.CENTER
    wcols = ["Họ", "Cấu hình thắng", "Precision @ recall 0.95", "Latency (ms)", "Flash (KB)", "Recall triển khai"]
    for j, name in enumerate(wcols):
        set_cell(wt.rows[0].cells[j], name, size=9, bold=True)
    style_header_row(wt.rows[0])
    for fam in fam_order:
        if fam not in within:
            continue
        r = by_key[(fam, within[fam])]
        cells = wt.add_row().cells
        vals = [FAM_LABEL[fam], within[fam], pm(r),
                fnum(r["latency_ms"], 1), fnum(r["flash_kb"], 1), fnum(r["recall_deploy"], 3)]
        is_overall = (fam == cfam)
        for j, v in enumerate(vals):
            set_cell(cells[j], v, size=9, bold=is_overall, fill="DDEBF7" if is_overall else None)
    para(doc, "Hàng xanh = mô hình thắng chung toàn cục (cross-family). “Recall triển khai” "
              "là recall trên DS2 khi áp ngưỡng tinh chỉnh từ DS1 — đo độ chuyển giao giữa bệnh "
              "nhân.", italic=True, size=9.5, color=GREY, space_after=8)

    # ===== 5. Phát hiện =====
    heading(doc, "5. Ba phát hiện chính", 1)
    bullet(doc, "đỉnh chất lượng nằm ở các model nhỏ; các cấu hình lớn nhất (cnn_c16-32-64-64, "
                "lstm_h32x2) vừa kém chất lượng vừa vượt 100 ms.", "1. “To hơn” không tốt hơn: ")
    bullet(doc, f"ngưỡng tinh chỉnh trên DS1 chỉ cho recall {rec_classic} trên DS2 với RF/SVM "
                f"(dịch chuyển giữa bệnh nhân), trong khi CNN/LSTM chuyển giao tốt ({rec_nn}). Đây "
                "là lý do nên ưu tiên mạng học trực tiếp từ nhịp thô.", "2. Khe hở hiệu chỉnh ngưỡng: ")
    bullet(doc, f"tại recall 0.95, precision trung bình của các họ chỉ {pmin:.2f}–{pmax:.2f} và "
                f"chồng lấn trong dải nhiễu seed (~±0.05) — không họ nào áp đảo; winner chọn theo "
                "trung bình. Phản ánh độ khó thực của bài toán mất cân bằng (~11% bất thường) với "
                "tín hiệu một chuyển đạo, chứ không phải mô hình kém.", "3. Độ khó thật: ")

    # ===== 6. Vì sao precision thấp =====
    heading(doc, "6. Vì sao precision thấp — rà soát đầy đủ", 1)
    para(
        doc,
        f"Tại điểm vận hành recall 0.95, precision của các họ chỉ {pmin:.2f}–{pmax:.2f}. Con số "
        "này phải đọc TƯƠNG ĐỐI: với lớp bất thường chỉ ~11%, một bộ phân loại đoán bừa theo tỷ "
        f"lệ đã cho precision ~0.11, nên {pmax:.2f} là gấp ~2 lần ngẫu nhiên. Quan trọng hơn, "
        "precision thấp KHÔNG phải vì mô hình kém — nó là tổng hợp của nhiều nguyên nhân, trong "
        "đó cần tách bạch phần “bản chất bài toán” (không bỏ được nếu muốn trung thực) với phần "
        "“do ta chọn” (còn dư địa cải thiện).",
    )

    # --- 6.1 bảng nguyên nhân: bản chất vs lựa chọn ---
    heading(doc, "6.1. Bản chất bài toán vs lựa chọn của ta", 2)
    ct = doc.add_table(rows=1, cols=3)
    ct.style = "Table Grid"; ct.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, name in enumerate(["Nguồn kéo precision xuống", "Bản chất hay lựa chọn?", "Bỏ được không?"]):
        set_cell(ct.rows[0].cells[j], name, size=9, bold=True)
    style_header_row(ct.rows[0])
    cause_rows = [
        ("Mất cân bằng lớp (~11% bất thường) → trần precision bị chặn bằng TOÁN HỌC",
         "Bản chất (tỷ lệ thật ngoài đời)", "Không"),
        ("Chia theo bệnh nhân (inter-patient, de Chazal) — test là người chưa từng thấy",
         "Bản chất (trung thực)", "Không — bỏ là ăn gian (rò rỉ)"),
        ("Hai lớp chồng lấn về SINH HỌC: nhịp SVEB trông gần y hệt nhịp thường",
         "Bản chất", "Không"),
        ("Ép recall ≥ 0.95 (điểm vận hành hung hãn để không bỏ sót)",
         "Lựa chọn (lý do lâm sàng)", "Có, nhưng đánh đổi recall"),
        ("Gộp nhị phân: SVEB (khó) lẫn VEB (dễ) chung một rổ ‘bất thường’",
         "Một phần lựa chọn, một phần bản chất", "Một phần"),
        ("Model tí hon + 1 chuyển đạo MLII (ràng buộc ESP32, vứt kênh V1)",
         "Lựa chọn (ràng buộc phần cứng)", "Có, nhưng phá ràng buộc ESP32"),
    ]
    for a, b, c in cause_rows:
        cells = ct.add_row().cells
        for j, v in enumerate((a, b, c)):
            set_cell(cells[j], v, size=8.5)
    para(doc, "Bốn nguyên nhân “bản chất” khiến precision thấp đồng đều ở MỌI họ — đó là dấu "
              "hiệu của trần khó bài toán, không phải lỗi một thuật toán. Hai nguyên nhân cuối là "
              "dư địa cải thiện thật sự (mục 6.3).", italic=True, size=9.5, color=GREY, space_after=10)

    # --- 6.2 SVEB vs VEB + nghịch lý AAMI ---
    heading(doc, "6.2. Thủ phạm lớn nhất: nhịp SVEB và nghịch lý nhãn AAMI", 2)
    para(doc,
         "Nhãn nhị phân của ta gom mọi nhịp bất thường vào một rổ, nhưng bên trong rổ đó có hai "
         "nhóm khó-dễ trái ngược, phân biệt theo NƠI xung điện phát ra:")
    bullet(doc, "xung phát thẳng từ tâm thất (buồng dưới), đi sai đường dẫn → tâm thất co bóp "
                "lệch → sóng QRS TO BÈ, MÉO MÓ, khác hẳn nhịp thường. Nhìn hình là biết ngay → "
                "mô hình bắt rất tốt (nhiều nghiên cứu đạt ~0.9/0.9). Ký hiệu gốc V, E.",
           "VEB (ngoại tâm thu thất) — DỄ: ")
    bullet(doc, "xung phát từ TRÊN tâm thất (tâm nhĩ hoặc nút AV) rồi vẫn chạy xuống theo đúng "
                "đường dẫn bình thường → tâm thất co bóp gần như bình thường → QRS TRÔNG GẦN Y "
                "HỆT nhịp Normal. Dấu hiệu phân biệt chỉ là nhịp đến sớm và sóng P hơi khác — mà "
                "sóng P lại nhìn rõ nhất ở kênh V1 (kênh ta đã vứt đi). Ký hiệu gốc A, a, J, S.",
           "SVEB (ngoại tâm thu trên thất) — KHÓ: ")
    add_image(doc, "beat_easy_vs_hard.png",
              "Hình 4. Nhịp thật từ MIT-BIH, so trong CÙNG một bệnh nhân. Trái: VEB (đỏ) đảo "
              "chiều, bè rộng, khác hẳn nhịp thường (xanh) → dễ. Phải: SVEB (cam) gần như chồng "
              "khít nhịp thường → gần như không phân biệt nổi bằng hình dạng.")
    para(doc,
         "Để đạt recall 0.95 trên CẢ rổ, mô hình buộc phải bắt cho được đám SVEB gần-giống-Normal "
         "(panel phải) → phải hạ ngưỡng rất thấp → quét nhầm hàng loạt nhịp Normal → precision "
         "sụp. SVEB là bài toán mở mà cả ngành đều bí trong điều kiện inter-patient 1-lead; nó bị "
         "gộp chung nên kéo điểm của VEB (vốn rất tốt) xuống theo.")
    para(doc,
         "Cộng thêm một nghịch lý của chuẩn gán nhãn AAMI: hình dạng và nhãn KHÔNG khớp gọn gàng.",
         space_after=2)
    at = doc.add_table(rows=1, cols=4)
    at.style = "Table Grid"; at.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, name in enumerate(["Siêu lớp AAMI", "Ký hiệu gốc", "Nhãn nhị phân của ta", "Hình dạng QRS"]):
        set_cell(at.rows[0].cells[j], name, size=9, bold=True)
    style_header_row(at.rows[0])
    aami_rows = [
        ("N (Normal)", "N, L, R, e, j", "0 — Bình thường", "đa số chuẩn; NHƯNG L/R (block nhánh) lại MÉO TO"),
        ("SVEB", "A, a, J, S", "1 — Bất thường", "gần y hệt Normal → khó"),
        ("VEB", "V, E", "1 — Bất thường", "méo bè, khác hẳn → dễ"),
        ("F (Fusion)", "F", "1 — Bất thường", "nhịp lai"),
        ("Q (Unknown)", "/, f, Q", "1 — Bất thường", "tạo nhịp / không phân loại (hiếm)"),
    ]
    for r0 in aami_rows:
        cells = at.add_row().cells
        for j, v in enumerate(r0):
            set_cell(cells[j], v, size=8.5)
    add_image(doc, "beat_morphology_gallery.png",
              "Hình 5. Một nhịp đại diện (thật, có chú thích record + vị trí mẫu) cho mỗi nhóm. "
              "N nhọn gọn; V (VEB) bè méo, đảo chiều; A (SVEB) gần y hệt N; L/R (block nhánh) "
              "méo rộng/notch nhưng AAMI vẫn gán nhãn 0 — Normal.")
    para(doc,
         "Nghịch lý: L (block nhánh trái) và R (block nhánh phải) có QRS méo to trông “bất "
         "thường”, nhưng AAMI vẫn xếp vào nhóm Normal (vì NGUỒN phát nhịp là xoang bình thường, "
         "chỉ đường dẫn trong thất bị nghẽn). Ngược lại, A/S (SVEB) trông bình thường nhưng phải "
         "gọi là bất thường. Mô hình thấy nhịp-méo-mà-gán-Normal lẫn nhịp-đẹp-mà-gán-bất-thường "
         "→ ranh giới hình dạng nhòe → precision khó lên. (Có thể đổi L/R sang ‘bất thường’ để số "
         "đẹp hơn, nhưng đó là đổi sang bài toán DỄ HƠN — không so được với chuẩn AAMI, nên ta "
         "giữ nguyên.)", italic=True, size=9.5, color=GREY, space_after=10)

    # --- 6.3 hướng xử lý ---
    heading(doc, "6.3. Hướng xử lý trong thực tế", 2)
    para(doc, "Phần dư địa cải thiện hợp lệ (không ăn gian), xếp theo mức tác động:", space_after=2)
    bullet(doc, "không báo động trên một nhịp lẻ mà yêu cầu N/M nhịp bất thường liên tiếp mới "
                "kích hoạt. Loạn nhịp thật thường kéo dài nhiều nhịp, nên cách này lọc bớt báo "
                "nhầm lẻ tẻ → nâng precision ở mức “sự kiện” rõ rệt mà gần như không giảm độ nhạy "
                "lâm sàng. Đòn bẩy mạnh và rẻ nhất.", "1. Gộp cảnh báo theo cửa sổ (debounce): ")
    bullet(doc, "ghi vài chục nhịp bình thường của CHÍNH bệnh nhân làm chuẩn rồi tinh chỉnh "
                "ngưỡng/mô hình (patient-adaptive — Hu 1997, de Chazal 2006, Kiranyaz 2016). Đây "
                "là cách mạnh nhất để thu hẹp khe hở hiệu chỉnh, nhưng đổi định nghĩa bài toán "
                "(không còn “đeo vào là chạy”).", "2. Cá nhân hóa theo bệnh nhân: ")
    bullet(doc, "thêm kênh V1 (nơi sóng P hiện rõ) là manh mối trực tiếp để gỡ SVEB — đúng "
                "điều các nghiên cứu inter-patient mạnh nhất làm (de Chazal dùng cả 2 lead). Đổi "
                "lại phải bỏ ràng buộc 1-cảm-biến của ESP32.", "3. Thêm chuyển đạo thứ hai (V1): ")
    bullet(doc, "thiết bị giữ recall cao (không bỏ sót), ca nghi ngờ đẩy lên bác sĩ / mô hình "
                "lớn ở cloud xác nhận → precision được “phục hồi” ở tầng sau mà không hy sinh an "
                "toàn ở thiết bị.", "4. Lọc hai tầng: ")
    bullet(doc, "báo cáo precision/recall riêng cho VEB và SVEB thay vì gộp — sẽ thấy VEB đẹp, "
                "SVEB xấu (đúng hiện trạng cả ngành), phơi bày đúng độ khó thật mà không đụng nhãn "
                "hay thổi phồng số.", "5. Báo cáo theo từng lớp: ")
    para(
        doc,
        "Tóm lại, với bài toán an toàn tim mạch thì giữ recall cao và chấp nhận precision thấp ở "
        "tầng thiết bị là lựa chọn đúng; precision được bù lại bằng lớp gộp cảnh báo và xác nhận "
        "phía sau. Phần precision thấp còn lại chủ yếu là TRẦN KHÓ trung thực (prevalence + SVEB "
        "+ inter-patient + 1-lead), không phải làm ẩu.",
        italic=True, size=10.5, color=GREY,
    )

    # ===== 7. Triển khai ESP32 =====
    heading(doc, "7. Triển khai và kiểm chứng trên ESP32", 1)
    para(doc, f"Cả {n_total} mô hình được convert sang nhân C thuần (không dùng TFLite Micro) và "
              "nạp chung một firmware trên ESP32-WROOM-32 (240 MHz, 320 KB RAM, 4 MB flash). Kết "
              "quả đo thật:")
    bullet(doc, f"{n_total}/{n_total} mô hình khớp nhãn bit-for-bit với PC — đảm bảo số liệu chất "
                "lượng trên PC là chuẩn tin cậy.", "Parity = 1.0000: ")
    bullet(doc, f"RAM làm việc ~5.5 KB, flash ~98% phân vùng huge_app cho cả {n_total} model — nút "
                "thắt là tính toán, không phải bộ nhớ.", "Tài nguyên: ")
    bullet(doc, "front-end db4 wavelet + FFT đo được 937 µs/nhịp, đã cộng vào độ trễ RF/SVM để so "
                "sánh công bằng.", "Trích đặc trưng: ")
    para(doc, "Lưu ý đo độ trễ: CNN/LSTM có kiến trúc cố định nên độ trễ/flash đo trên thiết bị "
              "không đổi sau khi làm sạch dữ liệu (tái dùng log benchmark). RF/SVM phụ thuộc dữ "
              "liệu (số node cây / số support vector) nên flash đã được tính lại; độ trễ đo lại là "
              "xấp xỉ nhưng đều ≪ 100 ms, không đổi kết luận khả thi.",
         italic=True, size=10, color=GREY)

    # ===== 8. Kết luận =====
    heading(doc, "8. Kết luận và hướng phát triển", 1)
    para(doc, f"Đã hoàn tất toàn bộ pipeline hai giai đoạn: huấn luyện + so sánh trên PC và triển "
              f"khai + đo thật trên ESP32. Mô hình đề xuất triển khai là {FAM_LABEL[cfam]}/{csize}: "
              f"ít báo động giả nhất trong các model hợp lệ, độ trễ {clat:.1f} ms, và ngưỡng chuyển "
              f"giao tốt giữa các bệnh nhân (recall {crecdep:.3f}).")
    para(doc, "Hướng phát triển tiếp theo:", space_after=2)
    bullet(doc, "tích hợp front-end trích đặc trưng vào đường suy luận RF/SVM ngay trên thiết bị.")
    bullet(doc, "lượng tử hóa int8 để giảm flash/độ trễ (hiện đang fp32).")
    bullet(doc, "thu hẹp khe hở hiệu chỉnh ngưỡng giữa các bệnh nhân cho RF/SVM.")

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
